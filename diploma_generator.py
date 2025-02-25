from pathlib import Path
from typing import List, Union
import fitz  # PyMuPDF for PDF handling
from docx import Document  # python-docx for Word documents
from PIL import Image, ImageDraw, ImageFont  # Pillow for image handling
import os
import subprocess  # For PDF conversion
import logging
from concurrent.futures import ThreadPoolExecutor
import random

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DiplomaGenerator:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.jpg', '.jpeg', '.png']
        self.template_path = None
        self.template_format = None
        self.soffice_ports = list(range(8100, 8115))  # 15 ports for parallel processing
        
        # Remove AI initialization for now
        # self.text_detector = pipeline("object-detection", model="microsoft/layoutlm-base-uncased")

    def load_template(self, template_path: Union[str, Path]) -> None:
        """Load the diploma template in any supported format"""
        template_path = Path(template_path)
        if not template_path.suffix.lower() in self.supported_formats:
            raise ValueError(f"Unsupported file format. Supported formats: {self.supported_formats}")
        
        self.template_path = template_path
        self.template_format = template_path.suffix.lower()
        logger.info(f"Template loaded: {template_path}")

    def load_names(self, names_path: Union[str, Path]) -> List[str]:
        """Load names from a text file (one name per line)"""
        with open(names_path, 'r', encoding='utf-8') as f:
            names = [name.strip() for name in f.readlines() if name.strip()]
            logger.info(f"Loaded {len(names)} names from {names_path}")
            return names

    def detect_placeholder(self) -> str:
        """Use AI to detect potential name placeholder in the template"""
        # Implementation depends on the file format
        if self.template_format in ['.jpg', '.jpeg', '.png']:
            return self._detect_placeholder_image()
        elif self.template_format == '.pdf':
            return self._detect_placeholder_pdf()
        else:  # Word documents
            return self._detect_placeholder_word()

    def generate_diplomas(self, names: List[str], output_dir: Union[str, Path], 
                         placeholder: str, output_format: str = 'docx') -> List[Path]:
        """Generate individual diplomas and return list of generated file paths"""
        output_dir = Path(output_dir)
        output_dir.mkdir(exist_ok=True, parents=True)
        
        generated_files = []
        for name in names:
            try:
                output_path = output_dir / f"diploma_{name.replace(' ', '_')}.{output_format}"
                self._generate_single_diploma(name, placeholder, output_path)
                generated_files.append(output_path)
                logger.info(f"Generated diploma for {name}")
            except Exception as e:
                logger.error(f"Failed to generate diploma for {name}: {str(e)}")
                continue
            
        return generated_files

    def _detect_placeholder_image(self) -> str:
        """Detect placeholder in image formats using OCR and AI"""
        image = Image.open(self.template_path)
        # Use pytesseract for OCR
        text = pytesseract.image_to_string(image)
        # Use AI to identify likely placeholder
        # Implementation details here

    def _detect_placeholder_pdf(self) -> str:
        """Detect placeholder in PDF format"""
        # Implementation using PyMuPDF

    def _detect_placeholder_word(self) -> str:
        """Detect placeholder in Word documents"""
        # Implementation using python-docx

    def _generate_single_diploma(self, name: str, placeholder: str, output_path: Path) -> None:
        """Generate a single diploma based on the template format"""
        if self.template_format in ['.jpg', '.jpeg', '.png']:
            self._generate_from_image(name, placeholder, output_path)
        elif self.template_format == '.pdf':
            self._generate_from_pdf(name, placeholder, output_path)
        else:  # Word documents
            self._generate_from_word(name, placeholder, output_path)

    def _generate_from_image(self, name: str, placeholder: str, output_path: Path) -> None:
        """Generate diploma from image template"""
        with Image.open(self.template_path) as img:
            # Create a copy to work with
            new_img = img.copy()
            draw = ImageDraw.Draw(new_img)
            
            # Basic implementation - you might want to adjust font size and position
            try:
                font = ImageFont.truetype("arial.ttf", 30)
            except OSError:
                font = ImageFont.load_default()
                
            # Center the name (basic implementation)
            w, h = img.size
            text_bbox = draw.textbbox((0, 0), name, font=font)
            text_w = text_bbox[2] - text_bbox[0]
            text_h = text_bbox[3] - text_bbox[1]
            x = (w - text_w) / 2
            y = (h - text_h) / 2
            
            draw.text((x, y), name, fill='black', font=font)
            new_img.save(output_path)

    def _generate_from_pdf(self, name: str, placeholder: str, output_path: Path) -> None:
        """Generate diploma from PDF template"""
        # Open the source PDF
        doc = fitz.open(self.template_path)
        
        # Create a new PDF document
        new_doc = fitz.open()
        
        # Copy pages from source to new document
        new_doc.insert_pdf(doc)
        
        # Process each page
        for page in new_doc:
            # Get all text instances
            text_instances = page.search_for(placeholder)
            
            # Replace each instance of the placeholder
            for inst in text_instances:
                # First remove the old text
                page.draw_rect(inst, color=fitz.utils.getColor('white'), fill=fitz.utils.getColor('white'))
                
                # Insert the new name
                # Get the position from the found instance
                x = inst[0]  # x coordinate of the placeholder
                y = inst[1]  # y coordinate of the placeholder
                
                # Insert the new text at the same position
                page.insert_text((x, y), name, 
                               fontname="helv",  # Use a standard font
                               fontsize=12,      # You might need to adjust this
                               color=fitz.utils.getColor('black'))
        
        # Save the modified document
        new_doc.save(output_path)
        
        # Close both documents
        new_doc.close()
        doc.close()

    def _generate_from_word(self, name: str, placeholder: str, output_path: Path) -> None:
        """Generate diploma from Word template"""
        # Create a copy of the template
        doc = Document(self.template_path)
        
        # Replace placeholder in paragraphs
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                # Preserve the original formatting
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, name)
        
        # Also check tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, name)
        
        # Save the modified document
        doc.save(output_path) 

    def _get_soffice_port(self):
        """Get a random available port from the pool"""
        return random.choice(self.soffice_ports)

    def convert_to_pdf(self, docx_path: Union[str, Path], pdf_path: Union[str, Path]) -> None:
        """Convert a single Word document to PDF using soffice"""
        try:
            logger.info(f"Converting {docx_path} to PDF")
            port = self._get_soffice_port()
            
            # Try conversion with specific port
            result = subprocess.run(
                ['soffice', 
                 f'--accept=socket,host=127.0.0.1,port={port};urp;StarOffice.ServiceManager', 
                 '--headless', 
                 '--convert-to', 'pdf', 
                 '--outdir', str(pdf_path.parent), 
                 str(docx_path)],
                capture_output=True,
                text=True
            )
            
            if result.returncode != 0:
                logger.warning(f"Conversion failed on port {port}: {result.stderr}")
                raise ValueError(f"Conversion failed: {result.stderr}")
                
            logger.info(f"Successfully converted {docx_path} to PDF using port {port}")
            
        except Exception as e:
            error_msg = f"Error converting {docx_path} to PDF: {str(e)}"
            logger.error(error_msg)
            raise ValueError(error_msg)

    def batch_convert_to_pdf(self, docx_dir: Union[str, Path], pdf_dir: Union[str, Path]) -> List[Path]:
        """Convert all Word documents in a directory to PDFs using parallel processing"""
        docx_dir = Path(docx_dir)
        pdf_dir = Path(pdf_dir)
        pdf_dir.mkdir(exist_ok=True, parents=True)
        
        docx_files = list(docx_dir.glob('*.docx'))
        converted_files = []
        
        # Use ThreadPoolExecutor for parallel processing
        with ThreadPoolExecutor(max_workers=15) as executor:
            futures = []
            for docx_file in docx_files:
                pdf_file = pdf_dir / f"{docx_file.stem}.pdf"
                future = executor.submit(self.convert_to_pdf, docx_file, pdf_file)
                futures.append((future, pdf_file))
            
            # Collect results
            for future, pdf_file in futures:
                try:
                    future.result()  # This will raise any exceptions that occurred
                    converted_files.append(pdf_file)
                except Exception as e:
                    logger.error(f"Failed to convert {pdf_file}: {str(e)}")
                    continue
        
        if not converted_files:
            raise ValueError("No files were successfully converted")
            
        return converted_files 